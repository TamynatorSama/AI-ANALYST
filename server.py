from fastapi import FastAPI, WebSocket, UploadFile, File, WebSocketDisconnect
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
from datetime import datetime
from uuid import uuid4
import json
import os
import re
import io
import logging
import traceback

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("dataan")

from agent import workflow
from langgraph.types import Command
from langchain_core.messages import ToolMessage, HumanMessage

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

# ── File-backed session store ────────────────────────────────────────────────

SESSIONS_FILE = DATA_DIR / "sessions.json"


def _load_sessions():
    if SESSIONS_FILE.exists():
        try:
            return json.loads(SESSIONS_FILE.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _save_sessions():
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    SESSIONS_FILE.write_text(json.dumps(sessions, indent=2), encoding="utf-8")


sessions = _load_sessions()


def get_session(sid):
    return sessions.get(sid)


# ── Pages ────────────────────────────────────────────────────────────────────

@app.get("/")
async def serve_app():
    return FileResponse(BASE_DIR / "static" / "index.html")


# ── REST API ─────────────────────────────────────────────────────────────────

@app.post("/api/sessions")
async def create_session(user_id: str = ""):
    sid = str(uuid4())
    sessions[sid] = {
        "id": sid,
        "owner": user_id,
        "name": "New Analysis",
        "created_at": datetime.now().isoformat(),
        "status": "idle",
        "data_path": None,
    }
    _save_sessions()
    return sessions[sid]


@app.get("/api/sessions")
async def list_sessions(user_id: str = ""):
    if user_id:
        return [s for s in sessions.values() if s.get("owner") == user_id]
    return list(sessions.values())


@app.delete("/api/sessions/{session_id}")
async def delete_session(session_id: str):
    sessions.pop(session_id, None)
    _save_sessions()
    return {"ok": True}


@app.post("/api/sessions/{session_id}/upload")
async def upload_file(session_id: str, file: UploadFile = File(...)):
    s = get_session(session_id)
    if not s:
        return JSONResponse(status_code=404, content={"error": "Session not found"})

    upload_dir = DATA_DIR / session_id / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path = upload_dir / file.filename

    with open(file_path, "wb") as f:
        f.write(await file.read())

    s["data_path"] = str(file_path)
    s["name"] = file.filename
    _save_sessions()
    # Log the upload as a chat message
    _append_chat_log(session_id, {"type": "user", "content": f"Uploaded {file.filename}"})
    return {"filename": file.filename, "path": str(file_path)}


@app.get("/api/sessions/{session_id}/chat")
async def get_chat_history(session_id: str):
    """Return saved chat messages for a session."""
    log_path = DATA_DIR / session_id / "chat_log.json"
    if not log_path.exists():
        return []
    try:
        return json.loads(log_path.read_text(encoding="utf-8"))
    except Exception:
        return []


@app.get("/api/sessions/{session_id}/images")
async def list_images(session_id: str, round: int = 0):
    """List images. round=0 means all rounds."""
    return _collect_images(session_id, round_num=round)


@app.get("/api/sessions/{session_id}/images/{round_num}/{filename}")
async def serve_image(session_id: str, round_num: int, filename: str):
    fp = DATA_DIR / session_id / "outputs" / f"round_{round_num}" / filename
    if not fp.exists():
        return JSONResponse(status_code=404, content={"error": "Not found"})
    return FileResponse(fp)


@app.get("/api/sessions/{session_id}/summary")
async def get_summary(session_id: str):
    summary_path = DATA_DIR / session_id / "executive_summary.md"
    if not summary_path.exists():
        return JSONResponse(status_code=404, content={"error": "No summary generated yet"})

    summary_md = summary_path.read_text(encoding="utf-8")
    all_images = _collect_images(session_id, round_num=0)

    return {"markdown": summary_md, "images": all_images}


@app.get("/api/sessions/{session_id}/summary/download")
async def download_summary(session_id: str):
    summary_path = DATA_DIR / session_id / "executive_summary.md"
    if not summary_path.exists():
        return JSONResponse(status_code=404, content={"error": "No summary generated yet"})

    md_text = summary_path.read_text(encoding="utf-8")
    docx_bytes = _md_to_docx(md_text, session_id)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=dataan_summary_{session_id[:8]}.docx"},
    )


# ── WebSocket ────────────────────────────────────────────────────────────────

@app.websocket("/ws/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: str):
    await websocket.accept()
    config = {"configurable": {"thread_id": session_id}}

    try:
        while True:
            raw = await websocket.receive_text()
            msg = json.loads(raw)

            if msg["type"] == "start":
                s = get_session(session_id)
                log.info(f"[START] session={session_id}, session_exists={s is not None}, data_path={s.get('data_path') if s else None}")
                if not s or not s.get("data_path"):
                    await _safe_send(websocket, {"type": "error", "detail": "Upload a file first."}, session_id)
                    continue

                s["status"] = "running"
                _save_sessions()
                await _safe_send(websocket, {"type": "status", "phase": "understanding", "detail": "Starting analysis..."}, session_id)

                try:
                    log.info(f"[AGENT] Starting workflow for {s['data_path']}")
                    async for event in workflow.astream(
                        {"data_path": s["data_path"], "user_id": session_id},
                        config=config,
                    ):
                        ws_events = _parse_event(event)
                        log.info(f"[AGENT] Event nodes: {list(event.keys())}, ws_events: {len(ws_events)}")
                        for ws_ev in ws_events:
                            await _safe_send(websocket, ws_ev, session_id)

                    log.info(f"[AGENT] Stream finished, checking interrupt/complete")
                    await _check_interrupt_or_complete(websocket, config, session_id)

                except Exception as e:
                    log.error(f"[AGENT] Error: {e}\n{traceback.format_exc()}")
                    await _safe_send(websocket, {"type": "error", "detail": str(e)}, session_id)

            elif msg["type"] == "answer":
                # Log the user's answer
                _append_chat_log(session_id, {"type": "user", "content": msg.get("text", "")})
                await _safe_send(websocket, {"type": "status", "phase": "understanding", "detail": "Processing your response..."}, session_id)

                try:
                    async for event in workflow.astream(
                        Command(resume=str(msg.get("text", ""))),
                        config=config,
                    ):
                        for ws_ev in _parse_event(event):
                            await _safe_send(websocket, ws_ev, session_id)

                    await _check_interrupt_or_complete(websocket, config, session_id)

                except Exception as e:
                    await _safe_send(websocket, {"type": "error", "detail": str(e)}, session_id)

            elif msg["type"] == "generate_summary":
                await _safe_send(websocket, {"type": "status", "phase": "summarizing", "detail": "Generating executive summary..."}, session_id)

                try:
                    async for event in workflow.astream(
                        Command(resume="__summary__"),
                        config=config,
                    ):
                        for ws_ev in _parse_event(event):
                            await _safe_send(websocket, ws_ev, session_id)

                    # Summary is done — send it to the client
                    summary_path = DATA_DIR / session_id / "executive_summary.md"
                    if summary_path.exists():
                        summary_md = summary_path.read_text(encoding="utf-8")
                        all_images = _collect_images(session_id, round_num=0)
                        s = get_session(session_id)
                        if s:
                            s["status"] = "complete"
                        await _safe_send(websocket, {
                            "type": "summary_ready",
                            "markdown": summary_md,
                            "images": all_images
                        }, session_id)
                    else:
                        await _safe_send(websocket, {"type": "error", "detail": "Summary generation failed"}, session_id)

                except Exception as e:
                    log.error(f"[SUMMARY] Error: {e}\n{traceback.format_exc()}")
                    await _safe_send(websocket, {"type": "error", "detail": str(e)}, session_id)

    except (WebSocketDisconnect, RuntimeError, Exception) as e:
        log.info(f"[WS] Connection closed for session {session_id}: {type(e).__name__}")


# ── Helpers ──────────────────────────────────────────────────────────────────

# Message types worth saving to the chat log
_CHAT_LOG_TYPES = {"message", "tool_call", "status", "question", "round_complete", "summary_ready", "error", "complete"}


async def _safe_send(ws: WebSocket, data: dict, session_id: str | None = None):
    """Send JSON over WebSocket. Also persist chat-worthy events to disk."""
    try:
        await ws.send_json(data)
    except Exception:
        pass
    # Persist to chat log
    if session_id and data.get("type") in _CHAT_LOG_TYPES:
        _append_chat_log(session_id, data)


def _append_chat_log(session_id: str, entry: dict):
    """Append a message to the session's chat log file."""
    log_path = DATA_DIR / session_id / "chat_log.json"
    log_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        existing = json.loads(log_path.read_text(encoding="utf-8")) if log_path.exists() else []
    except Exception:
        existing = []
    existing.append(entry)
    log_path.write_text(json.dumps(existing, default=str), encoding="utf-8")


async def _check_interrupt_or_complete(ws: WebSocket, config: dict, session_id: str):
    state = await workflow.aget_state(config)
    if state.next:
        tasks = state.tasks
        if tasks and tasks[0].interrupts:
            iv = tasks[0].interrupts[0].value

            # Distinguish between ask_user interrupt and round_complete interrupt
            if isinstance(iv, dict) and iv.get("type") == "round_complete":
                round_num = iv.get("round", 1)
                images = _collect_images(session_id, round_num=round_num)
                await _safe_send(ws, {
                    "type": "round_complete",
                    "round": round_num,
                    "images": images,
                    "image_count": len(images)
                }, session_id)
            else:
                # ask_user interrupt
                q = iv.get("question", "Please provide input:") if isinstance(iv, dict) else str(iv)
                await _safe_send(ws, {"type": "question", "text": q}, session_id)
    else:
        # Graph fully ended (after SummaryModel)
        s = get_session(session_id)
        if s:
            s["status"] = "complete"
            _save_sessions()
        summary_path = DATA_DIR / session_id / "executive_summary.md"
        if summary_path.exists():
            summary_md = summary_path.read_text(encoding="utf-8")
            all_images = _collect_images(session_id, round_num=0)
            await _safe_send(ws, {"type": "summary_ready", "markdown": summary_md, "images": all_images}, session_id)
        else:
            await _safe_send(ws, {"type": "complete", "images": _collect_images(session_id, round_num=0)}, session_id)

def _md_to_docx(md_text: str, session_id: str) -> bytes:
    """Convert markdown summary to a .docx Word document with embedded images."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title = doc.add_heading("DATAAN — Executive Summary", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Collect all images for [See: filename.png] resolution
    image_map = {}  # filename → absolute path
    outputs_dir = DATA_DIR / session_id / "outputs"
    if outputs_dir.exists():
        for rd in outputs_dir.iterdir():
            if rd.is_dir() and rd.name.startswith("round_"):
                for f in rd.iterdir():
                    if f.suffix.lower() in (".png", ".jpg", ".jpeg"):
                        image_map[f.name] = str(f)

    for line in md_text.split("\n"):
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph("_" * 50)
        elif stripped.startswith("- "):
            # Bullet point
            text = stripped[2:]
            text = _insert_image_refs(doc, text, image_map)
            if text is not None:
                p = doc.add_paragraph(style="List Bullet")
                _add_rich_text(p, text)
        elif re.match(r"^\d+\.\s", stripped):
            # Numbered list
            text = re.sub(r"^\d+\.\s", "", stripped)
            text = _insert_image_refs(doc, text, image_map)
            if text is not None:
                p = doc.add_paragraph(style="List Number")
                _add_rich_text(p, text)
        else:
            # Check for [See: filename.png] and embed image
            see_match = re.search(r"\[See:\s*(.+?)\]", stripped)
            if see_match:
                fname = see_match.group(1).strip()
                # Add text before the image ref
                before = stripped[:see_match.start()].strip()
                if before:
                    p = doc.add_paragraph()
                    _add_rich_text(p, before)
                # Embed image
                if fname in image_map:
                    try:
                        doc.add_picture(image_map[fname], width=Inches(5.5))
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cap = doc.add_paragraph(fname.replace(".png", "").replace("_", " ").title())
                        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cap.runs[0].font.size = Pt(9)
                        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    except Exception:
                        doc.add_paragraph(f"[Chart: {fname}]")
                # Add text after the image ref
                after = stripped[see_match.end():].strip()
                if after:
                    p = doc.add_paragraph()
                    _add_rich_text(p, after)
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                _add_rich_text(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _insert_image_refs(doc, text: str, image_map: dict):
    """If text contains [See: x.png], embed the image and return None. Otherwise return text."""
    see_match = re.search(r"\[See:\s*(.+?)\]", text)
    if see_match:
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        fname = see_match.group(1).strip()
        clean = re.sub(r"\*?\[See:\s*.+?\]\*?", "", text).strip()
        if clean:
            p = doc.add_paragraph()
            _add_rich_text(p, clean)
        if fname in image_map:
            try:
                doc.add_picture(image_map[fname], width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                doc.add_paragraph(f"[Chart: {fname}]")
        return None
    return text


def _add_rich_text(paragraph, text: str):
    """Parse **bold** in text and add runs to paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)



def _collect_images(session_id: str, round_num: int = 0):
    """Collect images. round_num=0 means all rounds."""
    base = DATA_DIR / session_id / "outputs"
    if not base.exists():
        return []
    imgs = []
    if round_num > 0:
        rd = base / f"round_{round_num}"
        if rd.exists():
            for f in sorted(rd.iterdir()):
                if f.suffix.lower() in (".png", ".jpg", ".jpeg", ".svg"):
                    imgs.append({"filename": f.name, "round": round_num, "url": f"/api/sessions/{session_id}/images/{round_num}/{f.name}"})
    else:
        # All rounds
        for rd in sorted(base.iterdir()):
            if rd.is_dir() and rd.name.startswith("round_"):
                rn = int(rd.name.replace("round_", ""))
                for f in sorted(rd.iterdir()):
                    if f.suffix.lower() in (".png", ".jpg", ".jpeg", ".svg"):
                        imgs.append({"filename": f.name, "round": rn, "url": f"/api/sessions/{session_id}/images/{rn}/{f.name}"})
    return imgs


def _parse_event(event):
    """Convert LangGraph stream event → list of WS-friendly dicts."""
    out = []
    for node, data in event.items():
        if node == "__interrupt__":
            continue
        msgs = data.get("messages", [])
        for m in msgs:
            cls = type(m).__name__

            if cls == "HumanMessage":
                out.append({"type": "message", "role": "user", "content": m.content, "node": node})

            elif cls == "AIMessage":
                text = ""
                if isinstance(m.content, str):
                    text = m.content
                elif isinstance(m.content, list):
                    for block in m.content:
                        if isinstance(block, dict) and block.get("type") == "text":
                            text += block.get("text", "")

                if text.strip():
                    out.append({"type": "message", "role": "ai", "content": text, "node": node})

                if hasattr(m, "tool_calls") and m.tool_calls:
                    for tc in m.tool_calls:
                        out.append({"type": "tool_call", "name": tc["name"], "node": node})

                        phase_map = {
                            "get_data_snapshot": ("understanding", "Analyzing your dataset..."),
                            "ask_user": ("asking", "Waiting for your input..."),
                            "write_to_file": ("understanding", "Recording analysis plan..."),
                            "execute_code": ("executing", "Running analysis..."),
                        }
                        if tc["name"] in phase_map:
                            p, d = phase_map[tc["name"]]
                            out.append({"type": "status", "phase": p, "detail": d})

            elif cls == "ToolMessage":
                content = m.content or ""
                if len(content) > 500:
                    content = content[:500] + "…"
                out.append({"type": "tool_result", "name": getattr(m, "name", ""), "output": content, "node": node})

    return out